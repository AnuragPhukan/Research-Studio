import os
from datetime import datetime
from docx import Document
from dotenv import load_dotenv
from pydantic import BaseModel
from langchain.agents import create_agent
from langchain_mistralai import ChatMistralAI
from langchain_core.messages import AIMessage
from tools import save_to_txt, search_tool, wiki_tool

load_dotenv()


class ResearchResponse(BaseModel):
    topic: str
    report: str
    sources: list[str]
    tools_used: list[str]


SYSTEM_PROMPT = (
    "You are an expert research assistant. Your task is to gather information on a "
    "given topic, write a detailed report, and list the sources and tools used in your "
    "research. Ensure that your responses are as detailed as possible, accurate, and "
    "well-structured. The report must be at least 500 words. "
    "Write a cohesive report in your own words; do not paste raw snippets or include "
    "URLs inside the report. "
    "Structure the report with headings and multiple paragraphs, including at least "
    "Introduction, Literature Review, Analysis/Discussion, and Conclusion sections. "
    "Always use the wikipedia tool at least once and include at least one Wikipedia "
    "source in the sources list. Provide at least three distinct sources. "
    "If the user asks to save results, use the save_to_txt tool."
)

TOOLS = [search_tool, wiki_tool, save_to_txt]

LLM = ChatMistralAI(
    model="codestral-latest",
    base_url=os.getenv("MISTRAL_BASE_URL"),
)

AGENT = create_agent(
    model=LLM,
    tools=TOOLS,
    system_prompt=SYSTEM_PROMPT,
    response_format=ResearchResponse,
)


def run_research(query: str) -> ResearchResponse | str:
    messages = [{"role": "user", "content": query}]
    structured_response = None

    for _ in range(3):
        result = AGENT.invoke({"messages": messages})
        structured_response = result.get("structured_response")

        if structured_response is None:
            messages_out = result.get("messages", [])
            last = messages_out[-1] if messages_out else None
            if isinstance(last, AIMessage):
                structured_response = last.content

        if isinstance(structured_response, ResearchResponse):
            tools_used = {tool.lower() for tool in structured_response.tools_used}
            report_text = structured_response.report
            has_wikipedia = "wikipedia" in tools_used
            word_count = len(report_text.split())
            distinct_sources = {
                source.strip().lower()
                for source in structured_response.sources
                if source.strip()
            }
            looks_like_snippets = "..." in report_text or "http" in report_text
            if (
                has_wikipedia
                and not looks_like_snippets
                and word_count >= 500
                and len(distinct_sources) >= 3
            ):
                break

        messages.append(
            {
                "role": "user",
                "content": (
                "Please use the wikipedia tool at least once and include a Wikipedia "
                "URL in the sources list. Rewrite the report as a cohesive narrative "
                "without raw snippets, ellipses, or URLs in the report. Expand the "
                "report to at least 500 words with section headings and multiple "
                "paragraphs, and include at least three distinct sources."
            ),
        }
    )

    return structured_response if structured_response is not None else "No response."


def format_output_text(response: ResearchResponse | str) -> str:
    if isinstance(response, ResearchResponse):
        return (
            f"Topic: {response.topic}\n"
            f"Report: {response.report}\n"
            f"Sources:\n- " + "\n- ".join(response.sources) + "\n"
            f"Tools Used:\n- " + "\n- ".join(response.tools_used) + "\n"
        )
    return str(response)


def save_report_docx(response: ResearchResponse | str, filepath: str) -> str:
    document = Document()
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if isinstance(response, ResearchResponse):
        document.add_heading("Research Report", level=1)
        document.add_paragraph(f"Timestamp: {timestamp}")
        document.add_heading("Topic", level=2)
        document.add_paragraph(response.topic)
        document.add_heading("Report", level=2)
        document.add_paragraph(response.report)
        document.add_heading("References", level=2)
        for source in response.sources:
            document.add_paragraph(source, style="List Bullet")
    else:
        document.add_heading("Research Report", level=1)
        document.add_paragraph(f"Timestamp: {timestamp}")
        document.add_paragraph(str(response))

    document.save(filepath)
    return filepath
